#!/usr/bin/env python3
"""
Extract Arkansas auctioneer exam data from DOCX and PDF files.
"""

import json
import re
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader

# Topic categories for mapping
TOPICS = [
    "Auction Procedures",
    "Licensing and Registration",
    "Board and Governance",
    "Legal and Contracts",
    "Duties and Responsibilities",
    "Financial and Records",
    "Penalties and Discipline",
    "Prohibited Acts and Violations",
    "Ethics and Conduct",
    "Definitions and Terminology"
]

def determine_topic(question_text, answer_text="", explanation=""):
    """Determine topic based on question content."""
    text = (question_text + " " + answer_text + " " + explanation).lower()
    
    # Board and Governance
    if any(word in text for word in ["board", "commissioner", "executive director", "governor", "appoint"]):
        return "Board and Governance"
    
    # Licensing and Registration
    if any(word in text for word in ["license", "registration", "permit", "apply", "applicant", "qualify"]):
        return "Licensing and Registration"
    
    # Legal and Contracts
    if any(word in text for word in ["contract", "legal", "law", "act", "statute", "agreement", "legal requirement"]):
        return "Legal and Contracts"
    
    # Auction Procedures
    if any(word in text for word in ["conduct", "auction", "bid", "sale", "start", "procedure", "process", "seller", "buyer"]):
        return "Auction Procedures"
    
    # Duties and Responsibilities
    if any(word in text for word in ["duty", "responsible", "obligation", "shall", "must", "require"]):
        return "Duties and Responsibilities"
    
    # Financial and Records
    if any(word in text for word in ["account", "record", "financial", "money", "payment", "deposit", "fund", "balance"]):
        return "Financial and Records"
    
    # Prohibited Acts and Violations
    if any(word in text for word in ["prohibited", "cannot", "illegal", "unlawful", "fraud", "unlicensed", "without license"]):
        return "Prohibited Acts and Violations"
    
    # Penalties and Discipline
    if any(word in text for word in ["penalty", "fine", "violation", "suspend", "revoke", "discipline", "criminal"]):
        return "Penalties and Discipline"
    
    # Ethics and Conduct
    if any(word in text for word in ["ethics", "ethical", "conduct", "honest", "integrity", "disqualify"]):
        return "Ethics and Conduct"
    
    # Definitions and Terminology
    if any(word in text for word in ["definition", "means", "defined as", "shall mean"]):
        return "Definitions and Terminology"
    
    # Default
    return "Auction Procedures"

def extract_questions_from_docx(docx_path):
    """Extract questions from DOCX file."""
    doc = Document(docx_path)
    questions = []
    
    i = 0
    question_id = 1
    
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # Look for question (doesn't start with bullet point, letter, or special chars)
        if text and not text.startswith(('A.', 'B.', 'C.', 'D.', 'Correct', 'Explanation', '·')):
            question_text = text
            options = {}
            answer = None
            explanation = None
            
            i += 1
            
            # Get options (A, B, C, D)
            while i < len(doc.paragraphs) and len(options) < 4:
                para = doc.paragraphs[i]
                text = para.text.strip()
                
                if not text:
                    i += 1
                    continue
                
                # Remove bullet point if present
                text = re.sub(r'^·\s+', '', text)
                
                # Try to match option pattern
                match = re.match(r'^([A-D])\.\s*(.+)$', text)
                if match:
                    letter = match.group(1)
                    content = match.group(2)
                    options[letter] = content
                    i += 1
                else:
                    break
            
            # Get correct answer
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                match = re.search(r'Correct answer:\s*([A-D])', text, re.IGNORECASE)
                if match:
                    answer = match.group(1)
                    i += 1
            
            # Get explanation
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                if text.lower().startswith('explanation:'):
                    explanation = re.sub(r'^explanation:\s*', '', text, flags=re.IGNORECASE)
                    i += 1
            
            # Validate and add question
            if question_text and len(options) == 4 and answer:
                try:
                    answer_idx = ord(answer) - ord('A')
                    
                    # Determine topic
                    topic = determine_topic(question_text, 
                                          " ".join(options.values()) if options else "",
                                          explanation or "")
                    
                    questions.append({
                        "id": question_id,
                        "question": question_text,
                        "options": [options.get(chr(65+j), "") for j in range(4)],
                        "correctAnswer": answer_idx,
                        "explanation": explanation or "",
                        "topic": topic
                    })
                    question_id += 1
                except Exception as e:
                    print(f"Error processing question: {e}")
                    i += 1
            else:
                i += 1
        else:
            i += 1
    
    return questions

def extract_study_guide_from_pdf(pdf_path):
    """Extract study guide content from PDF file."""
    reader = PdfReader(pdf_path)
    sections = []
    
    # Extract all text from PDF
    full_text = ""
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text += text + "\n"
    
    # Parse sections by looking for headings
    lines = full_text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            if current_content and current_section:
                # Save section if we have content
                if len(current_content) > 0:
                    sections.append({
                        "heading": current_section,
                        "content": [c for c in current_content if c]
                    })
                    current_section = None
                    current_content = []
            continue
        
        # Check if this is a potential heading (short line, mostly uppercase)
        word_count = len(stripped.split())
        uppercase_ratio = sum(1 for c in stripped if c.isupper()) / max(1, len(stripped.replace(' ', '')))
        
        is_heading = (word_count <= 10 and uppercase_ratio > 0.5)
        
        if is_heading and current_section is not None and current_content:
            # Save previous section
            sections.append({
                "heading": current_section,
                "content": [c for c in current_content if c]
            })
            current_section = stripped
            current_content = []
        elif is_heading:
            if current_section is None:
                current_section = stripped
            else:
                if current_content:
                    sections.append({
                        "heading": current_section,
                        "content": [c for c in current_content if c]
                    })
                current_section = stripped
                current_content = []
        else:
            if current_section:
                current_content.append(stripped)
    
    # Add last section
    if current_section and current_content:
        sections.append({
            "heading": current_section,
            "content": [c for c in current_content if c]
        })
    
    # Filter out very short sections
    sections = [s for s in sections if len(s.get("content", [])) > 0]
    
    return {
        "state": "Arkansas",
        "title": "Arkansas Auctioneer Study Guide",
        "sections": sections
    }

def main():
    """Main extraction function."""
    base_path = Path("/home/runner/work/Auction-Academy-Test-Prep/Auction-Academy-Test-Prep")
    
    # Extract questions
    print("Extracting questions from AR Question Bank.docx...")
    docx_file = base_path / "AR Question Bank.docx"
    questions = extract_questions_from_docx(docx_file)
    
    print(f"Extracted {len(questions)} questions")
    
    # Save questions
    output_questions = base_path / "src" / "data" / "ar_questions_temp.json"
    with open(output_questions, 'w') as f:
        json.dump(questions, f, indent=2)
    print(f"Saved to {output_questions}")
    
    # Extract study guide
    print("\nExtracting study guide from AR study guide.pdf...")
    pdf_file = base_path / "AR study guide.pdf"
    study_guide = extract_study_guide_from_pdf(pdf_file)
    
    print(f"Extracted {len(study_guide['sections'])} sections")
    
    # Save study guide
    output_guide = base_path / "src" / "data" / "ar_study_guide.json"
    with open(output_guide, 'w') as f:
        json.dump(study_guide, f, indent=2)
    print(f"Saved to {output_guide}")
    
    # Print summary
    print("\n" + "="*50)
    print("SUMMARY")
    print("="*50)
    print(f"Questions extracted: {len(questions)}")
    print("\nTopic distribution:")
    topic_counts = {}
    for q in questions:
        topic = q["topic"]
        topic_counts[topic] = topic_counts.get(topic, 0) + 1
    
    for topic, count in sorted(topic_counts.items()):
        print(f"  {topic}: {count}")
    
    return len(questions) >= 250  # Success if we got enough questions

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
